import shutil
import psutil

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponseForbidden
from django.views.generic import ListView, CreateView, UpdateView

from .mixins import RoleRequiredMixin
from .models import *
from django.contrib.auth import get_user_model
from django.urls import reverse_lazy
from .forms import *
from django.db.models import Q
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DetailView
from django.contrib.auth import logout

from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse
import os
import hashlib
from django.conf import settings
from django.http import FileResponse, Http404
import csv
from io import StringIO

# 如果要导出 Word/PDF，需要安装额外依赖：
# pip install python-docx reportlab
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

User = get_user_model()


@login_required
def dashboard(request):
    # ===== 系统资源部分 =====
    cpu = psutil.cpu_percent(interval=0.1)
    mem = psutil.virtual_memory()
    disk = shutil.disk_usage("/")

    disk_percent = round(disk.used / disk.total * 100, 1)

    # ===== 患者概览部分 =====
    patients = Patient.objects.all()
    total_patients = patients.count()

    # 每个卡片组对应的字段列表（来自 patient_form_partial.html）:contentReference[oaicite:0]{index=0}
    GROUP_FIELDS = {
        "basic": {
            "label": "一、基本信息",
            "fields": [
                "name", "gender", "handedness", "birthday",
                "department", "bed_number", "medical_record_number",
                "admission_date", "education_level", "occupation",
                "imaging_number", "admission_diagnosis",
            ],
        },
        "history": {
            "label": "二、病史",
            "fields": [
                "past_medical_history", "other_medical_history",
                "family_history", "first_seizure_age",
                "first_seizure_description", "medication_history",
            ],
        },
        "semiology": {
            "label": "三、发作症状学",
            "fields": [
                "aura", "typical_seizure_time", "typical_seizure_semiology",
                "seizure_duration_seconds", "seizure_duration_minutes",
                "seizure_freq_per_day", "seizure_freq_per_week",
                "seizure_freq_per_month", "seizure_freq_per_year",
            ],
        },
        "neuro": {
            "label": "四、神经系统检查",
            "fields": [
                "neuro_exam", "neuro_exam_description",
            ],
        },
        "cognitive": {
            "label": "五、认知和精神量表",
            "fields": [
                "moca_score", "hama_score", "hamd_score",
                "bai_score", "bdi_score", "epilepsy_scale_score",
            ],
        },
        "eeg": {
            "label": "六、视频头皮 EEG 检查",
            "fields": [
                "eeg_interictal", "eeg_ictal",
                "eeg_clinical_correlation", "eeg_file_link",
            ],
        },
        "imaging": {
            "label": "七、影像学检查",
            "fields": [
                "mri_brief", "mri_link",
                "pet_brief", "pet_link",
            ],
        },
        "first_stage": {
            "label": "八、一期无创性评估结果",
            "fields": [
                "first_stage_lateralization",
                "first_stage_region",
                "first_stage_location",
            ],
        },
        "seeg": {
            "label": "九、SEEG 发作间期及发作期放电",
            "fields": [
                "seeg_interictal_overall", "seeg_group1",
                "seeg_group2", "seeg_group3",
                "seeg_ictal", "seeg_file_link",
            ],
        },
        "second_stage": {
            "label": "十、二期有创性评估结果",
            "fields": [
                "second_stage_core_zone",
                "second_stage_hypothesis_zone",
            ],
        },
        "resection": {
            "label": "十一、外科切除计划",
            "fields": [
                "resection_plan_convex",
                "resection_plan_concave",
            ],
        },
        "evaluation": {
            "label": "十二、评估信息",
            "fields": [
                "evaluator", "evaluation_date",
            ],
        },
    }

    # 计算每个分组的完善情况
    group_stats = []
    if total_patients > 0:
        # 用 Python 循环，数量不大的情况下足够；以后患者量特别大再优化为 ORM 统计
        for key, cfg in GROUP_FIELDS.items():
            completed = 0
            for p in patients:
                # “全部填满” = 所有字段非空/非 None/非空串
                if all(getattr(p, field, None) not in (None, "", []) for field in cfg["fields"]):
                    completed += 1
            percent = round(completed / total_patients * 100, 1) if total_patients else 0.0
            group_stats.append({
                "key": key,
                "label": cfg["label"],
                "completed": completed,
                "total": total_patients,
                "percent": percent,
                "incomplete_percent": round(100 - percent, 1),
            })
    else:
        # 没有患者时，全部 0
        for key, cfg in GROUP_FIELDS.items():
            group_stats.append({
                "key": key,
                "label": cfg["label"],
                "completed": 0,
                "total": 0,
                "percent": 0.0,
                "incomplete_percent": 100.0,
            })

    context = {
        # 患者概览
        "total_patients": total_patients,
        "group_stats": group_stats,

        # 系统资源
        "cpu_percent": cpu,
        "cpu_free_percent": round(100 - cpu, 1),

        "mem_percent": mem.percent,
        "mem_free_percent": round(100 - mem.percent, 1),
        "mem_used": mem.used,
        "mem_total": mem.total,

        "disk_percent": disk_percent,
        "disk_free_percent": round(100 - disk_percent, 1),
        "disk_used": disk.used,
        "disk_total": disk.total,
    }
    return render(request, "epilepsy/dashboard.html", context)



class PatientListView(RoleRequiredMixin, ListView):
    model = Patient
    template_name = "epilepsy/patient_list.html"
    context_object_name = "patients"
    paginate_by = 20  # 每页 20 条，可按需要调整
    allowed_roles = [UserRole.ADMIN, UserRole.STAFF, UserRole.GUEST]

    def get_queryset(self):
        qs = super().get_queryset().order_by("id")
        q = self.request.GET.get("q", "").strip()
        if q:
            qs = qs.filter(
                Q(name__icontains=q)
                | Q(bed_number__icontains=q)
                | Q(department__icontains=q)
            )
        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["q"] = self.request.GET.get("q", "").strip()
        return context


class PatientCreateView(RoleRequiredMixin, CreateView):
    model = Patient
    form_class = PatientForm
    template_name = "epilepsy/patient_form.html"
    success_url = reverse_lazy("epilepsy:patient_list")
    allowed_roles = [UserRole.ADMIN, UserRole.STAFF]

    def form_valid(self, form):
        # 先按正常流程保存 Patient
        response = super().form_valid(form)

        # self.object 是刚保存好的 patient
        _handle_patient_file_uploads(self.request, self.object)

        # 如果是 AJAX（前端 XHR 设置了 X-Requested-With）
        if self.request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({
                "success": True,
                "patient_id": self.object.pk,
            })

        # 普通表单提交，继续走原来的重定向
        return response

    def form_invalid(self, form):
        # AJAX 情况下：返回 JSON，前端看到 success=False 就弹“保存失败”
        if self.request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({
                "success": False,
                "errors": form.errors,  # 你以后想在前端展示具体错误也可以用到
            })
        # 非 AJAX：保持原样，重新渲染模板
        return super().form_invalid(form)

class PatientUpdateView(RoleRequiredMixin, UpdateView):
    model = Patient
    form_class = PatientForm
    template_name = "epilepsy/patient_form_partial.html"  # for slide-in panel
    success_url = reverse_lazy("epilepsy:patient_list")
    allowed_roles = [UserRole.ADMIN, UserRole.STAFF]

    def get_template_names(self):
        # If AJAX request (for slide-in), use partial template
        if self.request.headers.get("x-requested-with") == "XMLHttpRequest":
            return ["epilepsy/patient_form_partial.html"]
        return ["epilepsy/patient_form.html"]


@login_required
def patient_delete(request, pk):
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role not in [UserRole.ADMIN, UserRole.STAFF]:
        return HttpResponseForbidden("无权限删除")

    patient = get_object_or_404(Patient, pk=pk)
    if request.method == "POST":
        patient.delete()
        return redirect("epilepsy:patient_list")
    return render(request, "epilepsy/patient_confirm_delete.html", {"patient": patient})

class UserListView(RoleRequiredMixin, ListView):
    model = User
    template_name = "epilepsy/user_list.html"
    context_object_name = "users"
    allowed_roles = [UserRole.ADMIN]


def _require_admin(request):
    """
    小工具函数：只有 ADMIN 才能操作用户。
    """
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role != UserRole.ADMIN:
        return HttpResponseForbidden("无权限操作")
    return None


def user_create(request):
    """
    新建用户（支持 AJAX 抽屉 & 普通页面）
    """
    deny = _require_admin(request)
    if deny:
        return deny

    if request.method == "POST":
        form = UserWithRoleForm(request.POST)
        if form.is_valid():
            user = form.save()
            # 抽屉 AJAX 提交
            if request.headers.get("x-requested-with") == "XMLHttpRequest":
                return JsonResponse({"success": True, "id": user.id})
            return redirect("epilepsy:user_list")
    else:
        form = UserWithRoleForm()

    # 抽屉里只要局部模板
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return render(request, "epilepsy/user_form_partial.html", {"form": form})

    # 兜底：完整页面编辑
    return render(request, "epilepsy/user_form.html", {"form": form})


def user_edit(request, pk):
    """
    编辑用户（支持 AJAX 抽屉 & 普通页面）
    """
    deny = _require_admin(request)
    if deny:
        return deny

    user_obj = get_object_or_404(User, pk=pk)

    if request.method == "POST":
        form = UserWithRoleForm(request.POST, instance=user_obj)
        if form.is_valid():
            form.save()
            if request.headers.get("x-requested-with") == "XMLHttpRequest":
                return JsonResponse({"success": True})
            return redirect("epilepsy:user_list")
    else:
        form = UserWithRoleForm(instance=user_obj)

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return render(request, "epilepsy/user_form_partial.html", {"form": form})

    return render(request, "epilepsy/user_form.html", {"form": form})


def user_toggle_active(request, pk):
    """
    启用 / 停用 用户
    """
    deny = _require_admin(request)
    if deny:
        return deny

    user_obj = get_object_or_404(User, pk=pk)

    if request.method != "POST":
        return HttpResponseForbidden("只允许 POST 请求")

    user_obj.is_active = not user_obj.is_active
    user_obj.save()

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse(
            {"success": True, "is_active": user_obj.is_active}
        )

    return redirect("epilepsy:user_list")


def user_delete(request, pk):
    """
    删除用户
    """
    deny = _require_admin(request)
    if deny:
        return deny

    user_obj = get_object_or_404(User, pk=pk)

    if request.method == "POST":
        user_obj.delete()
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": True})
        return redirect("epilepsy:user_list")

    # 简单确认页（可选）
    return render(
        request,
        "epilepsy/user_confirm_delete.html",
        {"user_obj": user_obj},
    )



class PatientDatasetListView(RoleRequiredMixin, DetailView):
    model = Patient
    template_name = "epilepsy/patient_datasets.html"
    context_object_name = "patient"
    allowed_roles = [UserRole.ADMIN, UserRole.STAFF, UserRole.GUEST]

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx["datasets"] = self.object.datasets.filter(is_active=True)
        return ctx


@login_required
def patient_dataset_download(request, pk):
    dataset = get_object_or_404(PatientDataset, pk=pk, is_active=True)
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role not in [
        UserRole.ADMIN,
        UserRole.STAFF,
        UserRole.GUEST,
    ]:
        return HttpResponseForbidden("无权限下载")

    # TODO: 在这里和 DGPF 集成
    # - 可以根据 dataset.globus_endpoint_id / dataset.globus_path
    #   构造一个链接跳转到 DGPF 的 transfer 页面
    # - 或者调用你 portal 中封装好的 transfer 视图

    return render(
        request,
        "epilepsy/download_not_implemented.html",
        {"dataset": dataset},
        status=501,  # 501 Not Implemented
    )

@login_required
def patient_files_panel(request, pk):
    """
    右侧抽屉用：显示某个患者的 MRI / PET / EEG / SEEG 文件列表。
    URL: /epilepsy/patients/<pk>/files/
    """
    patient = get_object_or_404(Patient, pk=pk)

    profile = getattr(request.user, "profile", None)
    if not profile or profile.role not in [
        UserRole.ADMIN,
        UserRole.STAFF,
        UserRole.GUEST,
    ]:
        return HttpResponseForbidden("无权限查看")

    context = {
        "patient": patient,
        "mri_files": patient.mri_files.all().order_by("-created_at"),
        "pet_files": patient.pet_files.all().order_by("-created_at"),
        "eeg_files": patient.eeg_files.all().order_by("-created_at"),
        "seeg_files": patient.seeg_files.all().order_by("-created_at"),
    }
    return render(request, "epilepsy/patient_files_panel.html", context)

def _get_display_value(obj, field_name):
    """
    优先用 get_xxx_display() 显示 choices 的中文，其次用原始值。
    """
    value = getattr(obj, field_name, "")
    if value is None:
        return ""
    method_name = f"get_{field_name}_display"
    if hasattr(obj, method_name):
        return getattr(obj, method_name)()
    return value


def _build_patient_sections(patient: Patient):
    """
    按照 “癫痫术前评估报告” 模板分区，组织所有字段。
    每个元素: (分区标题, [ (字段中文名, 值), ... ])
    """
    GROUP_FIELDS = [
        {
            "label": "一、基本信息",
            "fields": [
                "name", "gender", "handedness", "birthday",
                "department", "bed_number", "medical_record_number",
                "admission_date", "education_level", "occupation",
                "imaging_number", "admission_diagnosis",
            ],
        },
        {
            "label": "二、病史",
            "fields": [
                "past_medical_history", "other_medical_history",
                "family_history", "first_seizure_age",
                "first_seizure_description", "medication_history",
            ],
        },
        {
            "label": "三、发作症状学",
            "fields": [
                "aura", "typical_seizure_time", "typical_seizure_semiology",
                "seizure_duration_seconds", "seizure_duration_minutes",
                "seizure_freq_per_day", "seizure_freq_per_week",
                "seizure_freq_per_month", "seizure_freq_per_year",
            ],
        },
        {
            "label": "四、神经系统检查",
            "fields": [
                "neuro_exam", "neuro_exam_description",
            ],
        },
        {
            "label": "五、认知和精神量表",
            "fields": [
                "moca_score", "hama_score", "hamd_score",
                "bai_score", "bdi_score", "epilepsy_scale_score",
            ],
        },
        {
            "label": "六、视频头皮 EEG 检查",
            "fields": [
                "eeg_interictal", "eeg_ictal",
                "eeg_clinical_correlation", "eeg_file_link",
            ],
        },
        {
            "label": "七、影像学检查",
            "fields": [
                "mri_brief", "mri_link",
                "pet_brief", "pet_link",
            ],
        },
        {
            "label": "八、一期无创性评估结果",
            "fields": [
                "first_stage_lateralization",
                "first_stage_region",
                "first_stage_location",
            ],
        },
        {
            "label": "九、SEEG 发作间期及发作期放电",
            "fields": [
                "seeg_interictal_overall", "seeg_group1",
                "seeg_group2", "seeg_group3",
                "seeg_ictal", "seeg_file_link",
            ],
        },
        {
            "label": "十、二期有创性评估结果",
            "fields": [
                "second_stage_core_zone",
                "second_stage_hypothesis_zone",
            ],
        },
        {
            "label": "十一、外科切除计划",
            "fields": [
                "resection_plan_convex",
                "resection_plan_concave",
            ],
        },
        {
            "label": "十二、评估信息",
            "fields": [
                "evaluator", "evaluation_date",
            ],
        },
    ]

    sections = []
    for grp in GROUP_FIELDS:
        section_label = grp["label"]
        field_pairs = []
        for field_name in grp["fields"]:
            try:
                field_obj = patient._meta.get_field(field_name)
                label = field_obj.verbose_name or field_name
            except Exception:
                label = field_name
            display_value = _get_display_value(patient, field_name)
            field_pairs.append((str(label), "" if display_value is None else str(display_value)))
        sections.append((section_label, field_pairs))

    return sections

def patient_file_download(request, file_type, file_id):
    """
    下载 MRI / PET / EEG / SEEG 文件
    URL: /epilepsy/files/<file_type>/<id>/download/
    """
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role not in [
        UserRole.ADMIN,
        UserRole.STAFF,
        UserRole.GUEST,
    ]:
        return HttpResponseForbidden("无权限下载")

    model_map = {
        "mri": MRIFile,
        "pet": PETFile,
        "eeg": EEGFile,
        "seeg": SEEGFile,
    }

    model_cls = model_map.get(file_type)
    if model_cls is None:
        raise Http404("未知文件类型")

    file_obj = get_object_or_404(model_cls, pk=file_id)
    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")
    file_path = os.path.join(base_dir, file_obj.parent_path, file_obj.save_name)

    if not os.path.exists(file_path):
        raise Http404("文件不存在")

    return FileResponse(
        open(file_path, "rb"),
        as_attachment=True,
        filename=file_obj.file_name,
    )

def _handle_patient_file_uploads(request, patient):
    """
    统一处理 MRI / PET / EEG / SEEG 文件的上传和删除。

    - 上传字段名：mri_files, pet_files, eeg_files, seeg_files
    - 删除隐藏字段：delete_mri_file_ids, delete_pet_file_ids, delete_eeg_file_ids, delete_seeg_file_ids
    """
    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")

    # 各类型的配置
    config = {
        "mri": {
            "model": MRIFile,
            "input_name": "mri_files",
            "delete_field": "delete_mri_file_ids",
        },
        "pet": {
            "model": PETFile,
            "input_name": "pet_files",
            "delete_field": "delete_pet_file_ids",
        },
        "eeg": {
            "model": EEGFile,
            "input_name": "eeg_files",
            "delete_field": "delete_eeg_file_ids",
        },
        "seeg": {
            "model": SEEGFile,
            "input_name": "seeg_files",
            "delete_field": "delete_seeg_file_ids",
        },
    }

    for file_type, cfg in config.items():
        model_cls = cfg["model"]
        input_name = cfg["input_name"]
        delete_field = cfg["delete_field"]

        # 1) 删除已存在的文件（延迟到点击“保存”时统一删除）
        delete_ids_raw = request.POST.get(delete_field, "").strip()
        if delete_ids_raw:
            ids = [i for i in delete_ids_raw.split(",") if i]
            for file_obj in model_cls.objects.filter(id__in=ids, patient=patient):
                file_path = os.path.join(
                    base_dir,
                    file_obj.parent_path,
                    file_obj.save_name,
                )
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except OSError:
                        pass
                file_obj.delete()

        # 2) 处理新上传的文件列表
        files = request.FILES.getlist(input_name)
        if not files:
            continue

        # 每种类型一个子目录，例如 mri/1/, pet/1/, eeg/1/ 等
        parent_path = f"{file_type}/{patient.id}"
        abs_dir = os.path.join(base_dir, parent_path)
        os.makedirs(abs_dir, exist_ok=True)

        for uploaded in files:
            orig_name = uploaded.name
            md5 = hashlib.md5()
            sha256 = hashlib.sha256()

            # 一边写入磁盘一边计算 hash
            # 这里先用一个临时文件名，等 hash 算完后再决定最终名字
            temp_path = os.path.join(abs_dir, f"tmp_{uploaded.name}")
            with open(temp_path, "wb") as f:
                for chunk in uploaded.chunks():
                    md5.update(chunk)
                    sha256.update(chunk)
                    f.write(chunk)

            hash_code = md5.hexdigest()
            sha256_code = sha256.hexdigest()

            # 计算真正的保存文件名：hash_code + 扩展名
            _, ext = os.path.splitext(orig_name)
            ext = ext.lower()
            final_name = f"{hash_code}{ext}"
            final_path = os.path.join(abs_dir, final_name)

            # 重命名临时文件
            os.replace(temp_path, final_path)

            # 写数据库记录（save_name 由模型的 save() 自动生成）
            model_cls.objects.create(
                patient=patient,
                parent_path=parent_path,
                file_name=orig_name,
                hash_code=hash_code,
                sha256_code=sha256_code,
            )


def simple_logout(request):
    logout(request)
    return redirect("login")  # "login" 是 accounts/login/ 的 url name

def patient_edit(request, pk):
    patient = get_object_or_404(Patient, pk=pk)

    if request.method == "POST":
        form = PatientForm(request.POST, request.FILES, instance=patient)
        if form.is_valid():
            patient = form.save()
            _handle_patient_file_uploads(request, patient)

            # AJAX 保存成功返回 JSON
            if request.headers.get("x-requested-with") == "XMLHttpRequest":
                return JsonResponse({"success": True})
            # 非 AJAX 情况下，正常重定向回列表
            return redirect("epilepsy:patient_list")
    else:
        form = PatientForm(instance=patient)

    # AJAX 请求：仅返回 partial（不包 base 模板）
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return render(request, "epilepsy/patient_form_partial.html", {
            "form": form,
        })

    # 直接访问 /patients/<id>/edit/ 的 fallback，全页编辑也能用
    return render(request, "epilepsy/patient_edit.html", {
        "form": form,
        "patient": patient,
    })


def patient_detail(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    form = PatientForm(instance=patient)

    context = {
        "patient": patient,
        "form": form,
    }

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return render(request, "epilepsy/patient_detail_partial.html", context)

    return render(request, "epilepsy/patient_detail.html", context)

# def patient_detail(request, pk):
#     patient = get_object_or_404(Patient, pk=pk)

#     # 用和编辑时同一个 PatientForm，保证字段一致
#     form = PatientForm(instance=patient)

#     fields = []
#     for bound_field in form:   # 依次遍历表单字段
#         if bound_field.is_hidden:
#             continue

#         field = bound_field.field
#         raw_val = bound_field.value()
#         display_val = raw_val

#         # 处理 choices 字段，显示中文而不是内部代码
#         if getattr(field, "choices", None):
#             choices_dict = dict(field.choices)
#             if isinstance(raw_val, (list, tuple)):
#                 labels = [choices_dict.get(v, v) for v in raw_val]
#                 display_val = ", ".join(str(x) for x in labels if x not in [None, ""])
#             else:
#                 display_val = choices_dict.get(raw_val, raw_val)

#         # 简单格式化：None -> 空字符串
#         if display_val is None:
#             display_val = ""

#         fields.append({
#             "name": bound_field.name,
#             "label": bound_field.label,
#             "value": display_val,
#         })

#     context = {
#         "patient": patient,
#         "fields": fields,
#     }

#     # 抽屉里用的：AJAX 请求只要这一块 HTML
#     if request.headers.get("x-requested-with") == "XMLHttpRequest":
#         return render(request, "epilepsy/patient_detail_partial.html", context)

#     # 直接访问完整详情页（可选）
#     return render(request, "epilepsy/patient_detail.html", context)

@login_required
def patient_export(request, pk, fmt):
    """
    导出单个患者的所有信息为 CSV / Word / PDF：
    - 在 large_files/info/<patient_id>/ 下生成文件
    - 在 PatientInfoFile 中只保留该患者 + 该格式最新一条记录
    - 点击按钮时自动生成并直接下载
    """
    patient = get_object_or_404(Patient, pk=pk)

    # 权限与其它下载逻辑保持一致
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role not in [
        UserRole.ADMIN,
        UserRole.STAFF,
        UserRole.GUEST,
    ]:
        return HttpResponseForbidden("无权限导出")

    fmt = fmt.lower()
    if fmt == "csv":
        fmt_enum = PatientInfoFile.Format.CSV
        ext = "csv"
    elif fmt == "word":
        fmt_enum = PatientInfoFile.Format.WORD
        ext = "docx"
    elif fmt == "pdf":
        fmt_enum = PatientInfoFile.Format.PDF
        ext = "pdf"
    else:
        raise Http404("未知导出格式")

    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")
    parent_path = f"info/{patient.id}"
    abs_dir = os.path.join(base_dir, parent_path)
    os.makedirs(abs_dir, exist_ok=True)

    # 1) 删除该患者同一格式之前的导出文件及数据库记录
    old_qs = PatientInfoFile.objects.filter(patient=patient, format=fmt_enum)
    for old in old_qs:
        old_path = os.path.join(base_dir, old.parent_path, old.save_name)
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass
    old_qs.delete()

    # 2) 构造分区数据
    sections = _build_patient_sections(patient)

    # 3) 生成临时文件
    tmp_name = f"tmp_patient_info.{ext}"
    tmp_path = os.path.join(abs_dir, tmp_name)

    if fmt == "csv":
        # 简单格式：每行 [分区, 字段, 值]
        with open(tmp_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["分区", "字段", "值"])
            for section_title, fields in sections:
                writer.writerow([section_title, "", ""])
                for label, value in fields:
                    writer.writerow(["", label, value])

    elif fmt == "word":
        document = Document()
        document.add_heading("癫痫术前评估报告", level=1)
        document.add_paragraph(f"患者姓名：{patient.name}")
        document.add_paragraph(f"住院号：{patient.medical_record_number}")
        document.add_paragraph("")  # 空行

        for section_title, fields in sections:
            document.add_heading(section_title, level=2)
            table = document.add_table(rows=0, cols=2)
            for label, value in fields:
                row_cells = table.add_row().cells
                row_cells[0].text = str(label)
                row_cells[1].text = str(value or "")
            document.add_paragraph("")

        document.save(tmp_path)

    elif fmt == "pdf":
        c = canvas.Canvas(tmp_path, pagesize=A4)
        width, height = A4
        margin = 50
        y = height - margin

        # TODO：如需更好中文支持，可注册中文字体，如 STSong
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, "癫痫术前评估报告")
        y -= 24

        c.setFont("Helvetica", 10)
        header_lines = [
            f"患者姓名：{patient.name}",
            f"住院号：{patient.medical_record_number}",
        ]
        for line in header_lines:
            c.drawString(margin, y, line)
            y -= 14
        y -= 10

        def ensure_y():
            nonlocal y
            if y < margin:
                c.showPage()
                y = height - margin
                c.setFont("Helvetica", 10)

        for section_title, fields in sections:
            ensure_y()
            c.setFont("Helvetica-Bold", 11)
            c.drawString(margin, y, section_title)
            y -= 18

            c.setFont("Helvetica", 9)
            for label, value in fields:
                text = f"{label}：{value or ''}"
                # 简单切行：按长度截断
                line_len = 45
                lines = [text[i:i + line_len] for i in range(0, len(text), line_len)]
                for ln in lines:
                    ensure_y()
                    c.drawString(margin + 10, y, ln)
                    y -= 12
            y -= 8

        c.save()

    # 4) 计算 hash & sha256，并重命名为 hashcode.ext
    md5 = hashlib.md5()
    sha256 = hashlib.sha256()
    with open(tmp_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
            sha256.update(chunk)

    hash_code = md5.hexdigest()
    sha256_code = sha256.hexdigest()
    final_name = f"{hash_code}.{ext}"
    final_path = os.path.join(abs_dir, final_name)
    os.replace(tmp_path, final_path)

    # 5) 写数据库记录（file_name = 患者姓名_住院号）
    display_name_parts = [patient.name]
    if patient.medical_record_number:
        display_name_parts.append(patient.medical_record_number)
    display_name = "_".join(display_name_parts)

    info_file = PatientInfoFile.objects.create(
        patient=patient,
        format=fmt_enum,
        parent_path=parent_path,
        file_name=display_name,
        hash_code=hash_code,
        sha256_code=sha256_code,
    )

    # 6) 返回下载响应，文件名用 “患者姓名_住院号.xxx”
    download_filename = f"{display_name}.{ext}"
    return FileResponse(
        open(final_path, "rb"),
        as_attachment=True,
        filename=download_filename,
    )
