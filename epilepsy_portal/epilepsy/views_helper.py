# epilepsy/views_helper.py

import os
import csv
import shutil
import hashlib
import psutil
import zipfile
import tempfile
from pathlib import Path

from django.conf import settings
from django.http import HttpResponseForbidden
from django.shortcuts import get_object_or_404

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from .json import PATIENT_GROUP_FIELDS, FIELDS_FOR_EXPORT

from .models import (
    Patient,
    MRIFile, PETFile, EEGFile, SEEGFile,
    PatientInfoFile, UserRole,
)

# 这些字段是“逗号分隔存储”的多选 code，需要手动翻译成中文
MULTI_CHOICE_MAP = {
    "past_medical_history": dict(Patient.PAST_MEDICAL_HISTORY_CHOICES),
    "other_medical_history": dict(Patient.OTHER_MEDICAL_HISTORY_CHOICES),
    "eeg_interictal_state": dict(Patient.EEG_INTERICTAL_STATE_CHOICES),
    "eeg_interictal_location": dict(Patient.EEG_INTERICTAL_LOCATION_CHOICES),
    "eeg_interictal_morph": dict(Patient.EEG_INTERICTAL_MORPH_CHOICES),
    "eeg_interictal_amount": dict(Patient.EEG_INTERICTAL_AMOUNT_CHOICES),
    "eeg_interictal_pattern": dict(Patient.EEG_INTERICTAL_PATTERN_CHOICES),
    "eeg_interictal_eye_relation": dict(Patient.EEG_INTERICTAL_EYE_RELATED_CHOICES),
    "eeg_ictal_state": dict(Patient.EEG_INTERICTAL_STATE_CHOICES),
    "eeg_ictal_location": dict(Patient.EEG_INTERICTAL_LOCATION_CHOICES),
    "eeg_onset_pattern": dict(Patient.EEG_ONSET_PATTERN_CHOICES),
    "seeg_ictal_morph": dict(Patient.EEG_INTERICTAL_MORPH_CHOICES),
    "seeg_ictal_amount": dict(Patient.EEG_INTERICTAL_AMOUNT_CHOICES),
    "seeg_ictal_pattern": dict(Patient.EEG_INTERICTAL_PATTERN_CHOICES),
    "seeg_ictal_onset_pattern": dict(Patient.SEEG_ICTAL_ONSET_PATTERN_CHOICES),

}

# =======================
#  Dashboard 配置 & 计算
# =======================

def build_patient_group_stats(patients):
    """
    根据 PATIENT_GROUP_FIELDS 计算每个分组的完善情况。
    """
    total_patients = patients.count()
    group_stats = []

    if total_patients > 0:
        for key, cfg in PATIENT_GROUP_FIELDS.items():
            completed = 0
            for p in patients:
                if all(getattr(p, f, None) not in (None, "", []) for f in cfg["fields"]):
                    completed += 1
            percent = round(completed / total_patients * 100, 1)
            group_stats.append({
                "key": key,
                "label": cfg["label"],
                "completed": completed,
                "total": total_patients,
                "percent": percent,
                "incomplete_percent": round(100 - percent, 1),
            })
    else:
        for key, cfg in PATIENT_GROUP_FIELDS.items():
            group_stats.append({
                "key": key,
                "label": cfg["label"],
                "completed": 0,
                "total": 0,
                "percent": 0.0,
                "incomplete_percent": 100.0,
            })

    return total_patients, group_stats


def build_dashboard_context():
    """
    构造 dashboard 所需的 context（系统资源 + 患者概览）。
    """
    # 系统资源
    cpu = psutil.cpu_percent(interval=0.1)
    mem = psutil.virtual_memory()
    disk = shutil.disk_usage("/")

    disk_percent = round(disk.used / disk.total * 100, 1)

    # 患者概览
    patients = Patient.objects.all()
    total_patients, group_stats = build_patient_group_stats(patients)

    return {
        # 患者概览
        "total_patients": total_patients,
        "group_stats": group_stats,

        # 系统资源
        "cpu_percent": cpu,
        "cpu_free_percent": round(100 - cpu, 1),

        "mem_percent": mem.percent,
        "mem_free_percent": round(100 - mem.percent, 1),
        "mem_used": mem.used,
        "mem_total": mem.total,

        "disk_percent": disk_percent,
        "disk_free_percent": round(100 - disk_percent, 1),
        "disk_used": disk.used,
        "disk_total": disk.total,
    }


# =======================
#  权限 & 通用工具
# =======================

def require_admin(request):
    """
    只有 ADMIN 才能操作用户，否则返回 HttpResponseForbidden。
    """
    profile = getattr(request.user, "profile", None)
    if not profile or profile.role != UserRole.ADMIN:
        return HttpResponseForbidden("无权限操作")
    return None


def _get_field_choice_map(obj, field_name: str):
    """Return a {code: label} mapping for a model field's choices, if any.

    Uses Django's flatchoices when available (handles grouped choices).
    Returns None when the field does not exist or has no choices.
    """
    try:
        field = obj._meta.get_field(field_name)
    except Exception:
        return None

    choices = getattr(field, "flatchoices", None) or getattr(field, "choices", None)
    if not choices:
        return None

    try:
        return dict(choices)
    except TypeError:
        # Defensive: flatten grouped choices like [(group, [(k,v), ...]), ...]
        flat = []
        for item in choices:
            if isinstance(item, (list, tuple)) and len(item) == 2 and isinstance(item[1], (list, tuple)):
                flat.extend(item[1])
            else:
                flat.append(item)
        return dict(flat)


def get_display_value(obj, field_name):
    """
    1) Django 原生 choices（单选）：优先 get_xxx_display()
    2) choices + 多选值（逗号分隔字符串 / list / tuple / set）：自动用 choices 翻译
       - 兼容旧逻辑：如在 MULTI_CHOICE_MAP 中，优先使用手工映射
    3) ManyToMany（或类似 manager）：展示为 '，' 连接的对象字符串
    4) 其它字段：原样返回
    """
    value = getattr(obj, field_name, "")
    if value in (None, "", [], {}):
        return ""

    # 1) Django 原生 choices（单选）
    method_name = f"get_{field_name}_display"
    if hasattr(obj, method_name):
        try:
            disp = getattr(obj, method_name)()
            # 个别多选字段可能返回 list/tuple/set
            if isinstance(disp, (list, tuple, set)):
                return "，".join(str(x) for x in disp)
            return disp
        except Exception:
            # 个别自定义 Field 可能实现不标准，继续走下面兜底
            pass

    # 3) ManyToMany / 关联管理器：直接 join 关联对象的字符串
    if hasattr(value, "all") and callable(getattr(value, "all")):
        try:
            return "，".join(str(x) for x in value.all())
        except Exception:
            pass

    # 2) choices 多选（checkbox / 多选下拉等）
    mapping = MULTI_CHOICE_MAP.get(field_name) or _get_field_choice_map(obj, field_name)
    if mapping:
        # 逗号分隔字符串，例如 "A,B,C"
        if isinstance(value, str):
            codes = [v.strip() for v in value.split(",") if v.strip()]
            if len(codes) > 1:
                labels = [mapping.get(code, code) for code in codes]
                return "，".join(labels)
            if len(codes) == 1:
                return mapping.get(codes[0], codes[0])

        # list/tuple/set，例如 ["A", "B"]
        if isinstance(value, (list, tuple, set)):
            labels = [mapping.get(str(code), str(code)) for code in value]
            return "，".join(labels)

        # 单值兜底
        return mapping.get(value, value)

    # 4) 兜底
    return value

PATIENT_SECTION_FIELDS = [
    {
        "label": "一、基本信息",
        "fields": PATIENT_GROUP_FIELDS["basic"]["fields"],
    },
    {
        "label": "二、病史",
        "fields": PATIENT_GROUP_FIELDS["history"]["fields"],
    },
    {
        "label": "三、发作症状学",
        "fields": PATIENT_GROUP_FIELDS["semiology"]["fields"],
    },
    {
        "label": "四、神经系统检查",
        "fields": PATIENT_GROUP_FIELDS["neuro"]["fields"],
    },
    {
        "label": "五、认知和精神量表",
        "fields": PATIENT_GROUP_FIELDS["cognitive"]["fields"],
    },
    {
        "label": "六、视频头皮 EEG 检查",
        "fields": PATIENT_GROUP_FIELDS["eeg"]["fields"],
    },
    {
        "label": "七、影像学检查",
        "fields": PATIENT_GROUP_FIELDS["imaging"]["fields"],
    },
    {
        "label": "八、一期无创性评估结果",
        "fields": PATIENT_GROUP_FIELDS["first_stage"]["fields"],
    },
    {
        "label": "九、SEEG 发作间期及发作期放电",
        "fields": PATIENT_GROUP_FIELDS["seeg"]["fields"],
    },
    {
        "label": "十、二期有创性评估结果",
        "fields": PATIENT_GROUP_FIELDS["second_stage"]["fields"],
    },
    {
        "label": "十一、外科切除计划",
        "fields": PATIENT_GROUP_FIELDS["resection"]["fields"],
    },
    {
        "label": "十二、评估信息",
        "fields": PATIENT_GROUP_FIELDS["evaluation"]["fields"],
    },
]


def build_patient_sections(patient: Patient):
    """
    用于导出时按分区组织所有字段。
    返回: [(section_title, [(label, value), ...]), ...]
    """
    sections = []
    for grp in PATIENT_SECTION_FIELDS:
        section_label = grp["label"]
        field_pairs = []
        for field_name in grp["fields"]:
            try:
                field_obj = patient._meta.get_field(field_name)
                label = field_obj.verbose_name or field_name
            except Exception:
                label = field_name
            display_value = get_display_value(patient, field_name)
            field_pairs.append((str(label), "" if display_value is None else str(display_value)))
        sections.append((section_label, field_pairs))
    return sections


# =======================
#  文件上传 / 下载辅助
# =======================

def handle_patient_file_uploads(request, patient):
    """
    统一处理 MRI / PET / EEG / SEEG 文件的上传和删除。
    """
    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")

    config = {
        "mri": {
            "model": MRIFile,
            "input_name": "mri_files",
            "delete_field": "delete_mri_file_ids",
        },
        "pet": {
            "model": PETFile,
            "input_name": "pet_files",
            "delete_field": "delete_pet_file_ids",
        },
        "eeg": {
            "model": EEGFile,
            "input_name": "eeg_files",
            "delete_field": "delete_eeg_file_ids",
        },
        "seeg": {
            "model": SEEGFile,
            "input_name": "seeg_files",
            "delete_field": "delete_seeg_file_ids",
        },
    }

    for file_type, cfg in config.items():
        model_cls = cfg["model"]
        input_name = cfg["input_name"]
        delete_field = cfg["delete_field"]

        # 删除旧记录
        delete_ids_raw = request.POST.get(delete_field, "").strip()
        if delete_ids_raw:
            ids = [i for i in delete_ids_raw.split(",") if i]
            for file_obj in model_cls.objects.filter(id__in=ids, patient=patient):
                file_path = os.path.join(
                    base_dir,
                    file_obj.parent_path,
                    file_obj.save_name,
                )
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except OSError:
                        pass
                file_obj.delete()

        # 新上传文件（支持普通文件 + zip）
        uploads = request.FILES.getlist(input_name)
        if not uploads:
            continue

        parent_path = f"{file_type}/{patient.id}"
        abs_dir = os.path.join(base_dir, parent_path)
        os.makedirs(abs_dir, exist_ok=True)

        def _save_one_file_to_store(src_path: str, display_name: str):
            """
            将磁盘上的文件 src_path 写入 large_files，并写 DB。
            display_name 用于写入 file_name（展示给用户的名称）。
            """
            md5 = hashlib.md5()
            sha256 = hashlib.sha256()

            # 读入并计算 hash
            with open(src_path, "rb") as rf:
                for chunk in iter(lambda: rf.read(1024 * 1024), b""):
                    md5.update(chunk)
                    sha256.update(chunk)

            hash_code = md5.hexdigest()
            sha256_code = sha256.hexdigest()

            _, ext = os.path.splitext(display_name)
            ext = (ext or os.path.splitext(src_path)[1]).lower()
            final_name = f"{hash_code}{ext}"
            final_path = os.path.join(abs_dir, final_name)

            # 直接移动/覆盖（同 hash 会覆盖为相同内容）
            os.replace(src_path, final_path)

            model_cls.objects.create(
                patient=patient,
                parent_path=parent_path,
                file_name=display_name,
                hash_code=hash_code,
                sha256_code=sha256_code,
            )

        def _safe_extract_zip(zip_path: str, extract_dir: str):
            """
            防 Zip Slip：确保解压后的文件路径都在 extract_dir 内。
            """
            extract_root = Path(extract_dir).resolve()
            with zipfile.ZipFile(zip_path, "r") as zf:
                for member in zf.infolist():
                    # 跳过目录
                    if member.is_dir():
                        continue
                    target = (extract_root / member.filename).resolve()
                    if not str(target).startswith(str(extract_root) + os.sep):
                        raise ValueError(f"Unsafe zip entry path: {member.filename}")
                zf.extractall(extract_root)

        for uploaded in uploads:
            orig_name = uploaded.name or ""
            _, ext0 = os.path.splitext(orig_name)
            ext0 = (ext0 or "").lower()

            # 先把上传内容写到临时文件（同目录，避免跨盘移动问题）
            tmp_uploaded_path = os.path.join(abs_dir, f"tmp_upload_{orig_name}")
            with open(tmp_uploaded_path, "wb") as f:
                for chunk in uploaded.chunks():
                    f.write(chunk)

            is_zip = (ext0 == ".zip") or zipfile.is_zipfile(tmp_uploaded_path)

            if not is_zip:
                # 普通文件：直接入库
                _save_one_file_to_store(tmp_uploaded_path, orig_name)
                continue

            # zip：解压 -> 递归收集文件 -> 入库 -> 清理 zip & 解压目录
            extract_dir = tempfile.mkdtemp(prefix="tmp_extract_", dir=abs_dir)
            try:
                _safe_extract_zip(tmp_uploaded_path, extract_dir)

                # 删除原始 zip 临时文件
                try:
                    os.remove(tmp_uploaded_path)
                except OSError:
                    pass

                # 递归遍历解压目录
                for root, _, filenames in os.walk(extract_dir):
                    for fn in filenames:
                        # 可选：跳过 macOS 垃圾文件
                        if fn in (".DS_Store",) or fn.startswith("._"):
                            continue

                        full_path = os.path.join(root, fn)

                        # 用 zip 内相对路径作为展示名（可追溯来源）
                        rel_path = os.path.relpath(full_path, extract_dir)
                        display_name = f"{Path(orig_name).stem}/{rel_path}".replace("\\", "/")

                        # 为避免移动后破坏遍历，先把文件移动到 abs_dir 下的临时文件名
                        staging_path = os.path.join(abs_dir, f"tmp_zip_{hashlib.md5(display_name.encode('utf-8')).hexdigest()}")
                        os.replace(full_path, staging_path)

                        _save_one_file_to_store(staging_path, display_name)

            finally:
                # 清理解压目录
                try:
                    shutil.rmtree(extract_dir, ignore_errors=True)
                except OSError:
                    pass

def build_patient_file_path(model_cls, file_id):
    """
    公共的“根据模型和 id 找到物理文件路径”的逻辑。
    """
    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")
    file_obj = get_object_or_404(model_cls, pk=file_id)
    file_path = os.path.join(base_dir, file_obj.parent_path, file_obj.save_name)
    return file_obj, file_path


# =======================
#  导出辅助
# =======================

def generate_patient_info_file(patient, fmt: str):
    """
    生成患者信息文件并写入 PatientInfoFile 表。
    返回 (final_path, download_filename, fmt_enum)
    """
    fmt = fmt.lower()
    if fmt == "csv":
        fmt_enum = PatientInfoFile.Format.CSV
        ext = "csv"
    elif fmt == "word":
        fmt_enum = PatientInfoFile.Format.WORD
        ext = "docx"
    elif fmt == "pdf":
        fmt_enum = PatientInfoFile.Format.PDF
        ext = "pdf"
    else:
        raise ValueError("未知导出格式")

    base_dir = getattr(settings, "LARGE_FILE_BASE_DIR", settings.BASE_DIR / "large_files")
    parent_path = f"info/{patient.id}"
    abs_dir = os.path.join(base_dir, parent_path)
    os.makedirs(abs_dir, exist_ok=True)

    # 删除旧文件
    old_qs = PatientInfoFile.objects.filter(patient=patient, format=fmt_enum)
    for old in old_qs:
        old_path = os.path.join(base_dir, old.parent_path, old.save_name)
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass
    old_qs.delete()

    sections = build_patient_sections(patient)
    tmp_name = f"tmp_patient_info.{ext}"
    tmp_path = os.path.join(abs_dir, tmp_name)

    if fmt == "csv":
        with open(tmp_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["分区", "字段", "值"])
            for section_title, fields in sections:
                writer.writerow([section_title, "", ""])
                for label, value in fields:
                    writer.writerow(["", label, value])
    elif fmt == "word":
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

        document = Document()

        # 1) 全局：把 Normal 段落“段前/段后”压到 0，行距单倍
        normal = document.styles["Normal"]
        normal_pf = normal.paragraph_format
        normal_pf.space_before = Pt(0)
        normal_pf.space_after = Pt(0)
        normal_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        def tighten_paragraph(p, after=0, before=0):
            pf = p.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            return p

        # 2) 标题（Heading1 默认太松，手动压一下）
        h1 = document.add_heading("癫痫术前评估报告", level=1)
        tighten_paragraph(h1, after=4, before=0)

        p1 = document.add_paragraph(f"患者姓名：{patient.name}")
        tighten_paragraph(p1, after=0, before=0)

        p2 = document.add_paragraph(f"住院号：{patient.medical_record_number}")
        tighten_paragraph(p2, after=2, before=0)

        for section_title, fields in sections:
            # 3) section 标题压紧（Heading2 默认段前段后很大）
            h2 = document.add_heading(section_title, level=2)
            tighten_paragraph(h2, after=2, before=4)

            # 4) 表格：固定两列宽 + 单元格段落压紧
            table = document.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            label_w = Cm(5)
            value_w = Cm(11.5)

            for label, value in fields:
                row_cells = table.add_row().cells

                row_cells[0].width = label_w
                row_cells[1].width = value_w

                row_cells[0].text = str(label)
                row_cells[1].text = str(value or "")

                row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                # 压紧单元格内部段落（否则每行还是会高）
                for c in row_cells:
                    for p in c.paragraphs:
                        tighten_paragraph(p, after=0, before=0)

            # section 之间只留一点点间隔（不要用很多空行）
            spacer = document.add_paragraph("")
            tighten_paragraph(spacer, after=2, before=0)

        document.save(tmp_path)
    
    elif fmt == "pdf":
        c = canvas.Canvas(tmp_path, pagesize=A4)
        width, height = A4
        margin = 50
        y = height - margin

        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, "癫痫术前评估报告")
        y -= 24

        c.setFont("Helvetica", 10)
        header_lines = [
            f"患者姓名：{patient.name}",
            f"住院号：{patient.medical_record_number}",
        ]
        for line in header_lines:
            c.drawString(margin, y, line)
            y -= 14
        y -= 10

        def ensure_y():
            nonlocal y
            if y < margin:
                c.showPage()
                y = height - margin
                c.setFont("Helvetica", 10)

        for section_title, fields in sections:
            ensure_y()
            c.setFont("Helvetica-Bold", 11)
            c.drawString(margin, y, section_title)
            y -= 18

            c.setFont("Helvetica", 9)
            for label, value in fields:
                text = f"{label}：{value or ''}"
                line_len = 45
                lines = [text[i:i + line_len] for i in range(0, len(text), line_len)]
                for ln in lines:
                    ensure_y()
                    c.drawString(margin + 10, y, ln)
                    y -= 12
            y -= 8

        c.save()

    # hash & rename
    md5 = hashlib.md5()
    sha256 = hashlib.sha256()
    with open(tmp_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
            sha256.update(chunk)

    hash_code = md5.hexdigest()
    sha256_code = sha256.hexdigest()
    final_name = f"{hash_code}.{ext}"
    final_path = os.path.join(abs_dir, final_name)
    os.replace(tmp_path, final_path)

    # DB record
    display_name_parts = [patient.name]
    if patient.medical_record_number:
        display_name_parts.append(patient.medical_record_number)
    display_name = "_".join(display_name_parts)

    PatientInfoFile.objects.create(
        patient=patient,
        format=fmt_enum,
        parent_path=parent_path,
        file_name=display_name,
        hash_code=hash_code,
        sha256_code=sha256_code,
    )

    download_filename = f"{display_name}.{ext}"
    return final_path, download_filename, fmt_enum
